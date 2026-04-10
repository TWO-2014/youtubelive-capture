#!/usr/bin/env python3
"""DOCX + PDF レポート生成: 冒頭に全体構成、後半に詳細"""

import json
import os
import sys
import subprocess
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def set_landscape(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_body(doc, text, bold=False, size=9):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
                r.bold = True
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def main():
    output_dir = sys.argv[1] if len(sys.argv) > 1 else "output/6h_full"

    with open(os.path.join(output_dir, "claude_analysis.json"), "r", encoding="utf-8") as f:
        summary = json.load(f)
    with open(os.path.join(output_dir, "chunk_details.json"), "r", encoding="utf-8") as f:
        chunks = json.load(f)
    gemini = None
    gpath = os.path.join(output_dir, "gemini_visual.json")
    if os.path.exists(gpath):
        with open(gpath, "r", encoding="utf-8") as f:
            gemini = json.load(f)

    products = summary["products"]
    patterns = summary["pattern_analysis"]
    tech_dist = patterns["technique_distribution"]
    info = summary["broadcast_info"]
    tech_jp = {"scarcity": "希少性", "social_proof": "社会的証明", "anchoring": "アンカリング",
               "comparison": "比較", "authority": "権威", "urgency": "緊急性"}
    total_tech = sum(tech_dist.values())

    doc = Document()
    set_landscape(doc)

    # ============================================================
    # 表紙
    # ============================================================
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ショップチャンネル 6時間放送\n分析レポート")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("セールストーク構造 × 番組構成 × CV導線 の統合分析")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"分析日: {info['analysis_date']}",
        f"URL: {info['url']}",
        f"キャンペーン: {info['campaign']}",
        "分析手法: Claude Code テキスト分析 + Gemini 2.5 Flash 視覚分析",
    ]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ============================================================
    # 目次
    # ============================================================
    add_styled_heading(doc, "目次", level=1)
    toc_items = [
        "1. エグゼクティブサマリー",
        "2. 商品タイムライン",
        "3. セールストーク パターン分析",
        "    3.1 共通トーク構造",
        "    3.2 テクニック分布",
        "    3.3 頻出フレーズ TOP 10",
        "4. 番組構成・CV導線の設計分析",
        "    4.1 フェーズ別時間配分",
        "    4.2 CTA設計パターン",
        "5. インサイト・示唆（6つの応用知見）",
        "6. 【詳細】商品別分析（テクニック実例・引用付き）",
    ]
    for p_name in products:
        toc_items.append(f"    6.{products.index(p_name)+1} {p_name['name'][:35]}")
    toc_items.append("7. 【詳細】Gemini 視覚分析ハイライト")
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    doc.add_page_break()

    # ============================================================
    # 1. エグゼクティブサマリー
    # ============================================================
    add_styled_heading(doc, "1. エグゼクティブサマリー", level=1)

    # KPI行
    kpi_table = doc.add_table(rows=1, cols=6)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [
        ("7", "商品数"), ("360分", "放送時間"), (str(total_tech), "テクニック検出"),
        (f"{info['total_segments']:,}", "セグメント"), ("360", "スクリーンショット"), ("60", "視覚分析ポイント"),
    ]
    for i, (val, label) in enumerate(kpis):
        cell = kpi_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val + "\n")
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    add_body(doc, (
        "6時間の放送で7つの商品セグメントを分析。1商品あたり平均45-55分のサイクルで、"
        "「導入→実演デモ(複数ラウンド)→価格発表(アンカリング)→社会的証明→在庫カウントダウン"
        "(希少性エスカレーション)→最終CTA→クロージング」の定型構造を持つ。"
        f"最も多用されるテクニックは希少性（{tech_dist['scarcity']}件）で、リアルタイム在庫数カウントダウンが全セグメントで使われている。"
        "全員送料無料キャンペーン（4/10-16）が横断的な緊急性ドライバーとして機能。"
    ), size=10)

    # ============================================================
    # 2. 商品タイムライン
    # ============================================================
    add_styled_heading(doc, "2. 商品タイムライン", level=1)
    rows = []
    for i, p in enumerate(products):
        dur = p.get("total_duration_minutes", p.get("duration_minutes", ""))
        rows.append([
            p["name"][:32],
            f'{p["start_time"]}〜{p["end_time"]}',
            f"{dur}分",
            str(p.get("price", ""))[:45],
            p.get("category", ""),
            p.get("brand", "")[:20],
        ])
    add_table(doc, ["商品名", "時間帯", "時間", "価格", "カテゴリ", "ブランド"], rows,
              col_widths=[7, 3.5, 1.5, 6, 2.5, 4])

    # ============================================================
    # 3. セールストーク パターン分析
    # ============================================================
    doc.add_page_break()
    add_styled_heading(doc, "3. セールストーク パターン分析", level=1)

    add_styled_heading(doc, "3.1 共通トーク構造", level=2)
    add_body(doc, patterns["common_structure"], size=10)

    add_styled_heading(doc, "3.2 テクニック分布", level=2)
    tech_rows = []
    for k, v in sorted(tech_dist.items(), key=lambda x: -x[1]):
        pct = round(v / total_tech * 100, 1)
        bar = "█" * round(v / max(tech_dist.values()) * 20)
        tech_rows.append([tech_jp.get(k, k), str(v), f"{pct}%", bar])
    add_table(doc, ["テクニック", "検出数", "割合", "分布"], tech_rows, col_widths=[4, 2, 2, 8])

    add_styled_heading(doc, "3.3 頻出フレーズ TOP 10", level=2)
    phrase_rows = [[f"「{ph['phrase']}」", str(ph["count"]), ph.get("context", "")] for ph in patterns.get("top_phrases", [])]
    add_table(doc, ["フレーズ", "回数", "文脈"], phrase_rows, col_widths=[8, 1.5, 10])

    # ============================================================
    # 4. 番組構成・CV導線
    # ============================================================
    doc.add_page_break()
    add_styled_heading(doc, "4. 番組構成・CV導線の設計分析", level=1)

    add_styled_heading(doc, "4.1 フェーズ別時間配分（平均）", level=2)
    phase_jp = {"intro": "導入", "demo": "実演デモ", "price_reveal": "価格発表",
                "social_proof": "社会的証明", "cta_scarcity": "CTA・在庫訴求", "close": "クロージング"}
    phase_rows = [[phase_jp.get(k, k), f"{v}%", "█" * round(v / 5)] for k, v in patterns.get("avg_phase_duration_pct", {}).items()]
    add_table(doc, ["フェーズ", "時間比率", "分布"], phase_rows, col_widths=[5, 2, 10])

    add_styled_heading(doc, "4.2 CTA設計パターン", level=2)
    cta_patterns = [
        ("電話番号", "常時表示。「0120から始まるフリーダイヤル」「タッチでショップ2番」で注文切替。"),
        ("Web/QR", "「スマホ・パソコンでショップチャンネルと検索」「QRコード」。電話混雑時の代替チャネルとして案内。"),
        ("在庫カウントダウン", "リアルタイムで「○○点を切りました」を1セグメント10-25回更新。サイズ・色別に詳細報告。"),
        ("注文件数報告", "「お電話○○名」「○○件のオーダー」でバンドワゴン効果。最大1万件突破の報告。"),
        ("送料無料キャンペーン", "全セグメントで繰り返し（推定30回以上）。「今日からスタート」「16日まで」「別住所への送付も無料」。"),
        ("ショップチャンネルカード", "5%還元を主要セグメントで告知。"),
        ("マッチングプライス", "複数商品セット購入で割引（メルティリッチダウンで使用: かけ敷き12,960円）。"),
    ]
    add_table(doc, ["CTA種別", "設計パターン"], cta_patterns, col_widths=[4, 20])

    # ============================================================
    # 5. インサイト・示唆
    # ============================================================
    doc.add_page_break()
    add_styled_heading(doc, "5. インサイト・示唆 — 他の販売チャネルに応用可能な知見", level=1)
    insights = [
        ("在庫カウントダウンのリアルタイム性",
         f"「残り○○点」の連続更新が最強の購買ドライバー（{tech_dist['scarcity']}件検出）。1セグメントで10-25回更新される。ECサイトでも在庫数のリアルタイム表示、「残りN点」通知、sold-out表示が購買転換率を大幅に改善する。"),
        ("デモの反復構造（全体の45%）",
         "1商品につき同じデモを2-3ラウンド繰り返す設計。途中参加の視聴者に対応しつつ、繰り返しで購買意欲を高める二重効果。動画コマースでも「途中から見ても分かる」ループ構成が重要。"),
        ("アンカリングの二重構造",
         "「メーカー希望小売価格→ショップチャンネル価格」の空間的アンカリングに加え、「原材料高騰で本来値上げすべきところを据え置き」「30周年特別」という時間軸のアンカリングを組み合わせる。単純な値引きより「企業努力」のストーリーが効く。"),
        ("社会的証明のライブ感",
         f"注文件数・同時通話数をリアルタイムで報告（{tech_dist['social_proof']}件検出）。「お電話350名」「1万件突破」に加え、顧客メッセージの即時読み上げ。バンドワゴン効果とFOMOを同時に発動。レビュー数表示やリアルタイム購入通知に応用可能。"),
        ("ゲスト出演の権威効果",
         f"社長・デザイナー・実演販売士が全商品に登場（権威テクニック{tech_dist['authority']}件）。開発背景や職人のこだわりを「本人の口」から語ることで説得力が格段に増す。EC上でも創業者ストーリーやメーカー担当者インタビューの埋め込みが有効。"),
        ("送料無料を横断ドライバーに",
         "期間限定の送料無料キャンペーン（推定30回以上言及）が個別商品訴求とは別レイヤーで全セグメントを貫通。「今買う理由」を商品特性に依存せず生成する仕組み。ECでも「全品送料無料ウィーク」は個別クーポンより効果的な場合がある。"),
    ]
    for i, (title, body) in enumerate(insights, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}")
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(0x4e, 0x79, 0xa7)
        add_body(doc, body, size=9)
        doc.add_paragraph()  # spacing

    # ============================================================
    # 6. 商品別詳細分析
    # ============================================================
    doc.add_page_break()
    add_styled_heading(doc, "6. 【詳細】商品別分析", level=1)
    add_body(doc, "各商品のセールストークに使用された説得テクニックを、原文引用とタイムスタンプ付きで記録。", size=9)

    for pi, prod in enumerate(products):
        doc.add_page_break()
        ts = prod.get("techniques_summary", {})
        dur = prod.get("total_duration_minutes", prod.get("duration_minutes", "?"))

        add_styled_heading(doc, f"6.{pi+1} {prod['name']}", level=2)

        # 基本情報テーブル
        info_rows = [
            ["ブランド", str(prod.get("brand", "") or ""), "品番", str(prod.get("item_number", "N/A") or "N/A")],
            ["時間帯", f"{prod['start_time']}〜{prod['end_time']} ({dur}分)", "カテゴリ", str(prod.get("category", "") or "")],
            ["価格", str(prod.get("price", "") or ""), "定価", str(prod.get("retail_price", prod.get("discount", "")) or "")],
        ]
        info_table = doc.add_table(rows=3, cols=4)
        info_table.style = 'Light Grid Accent 1'
        for ri, row in enumerate(info_rows):
            for ci, val in enumerate(row):
                cell = info_table.rows[ri].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
                        if ci % 2 == 0:
                            r.bold = True

        # テクニック分布
        doc.add_paragraph()
        add_styled_heading(doc, "テクニック分布", level=3)
        tech_summary_rows = []
        for tk in ["scarcity", "social_proof", "anchoring", "comparison", "authority", "urgency"]:
            tv = ts.get(tk, 0)
            bar = "█" * round(tv / max(max(ts.values(), default=1), 1) * 15) if tv else ""
            tech_summary_rows.append([tech_jp[tk], str(tv), bar])
        add_table(doc, ["テクニック", "件数", "分布"], tech_summary_rows, col_widths=[4, 1.5, 8])

        # キーフレーズ
        add_styled_heading(doc, "キーフレーズ", level=3)
        for kp in prod.get("key_phrases", [])[:6]:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"「{kp}」")
            run.font.size = Pt(9)

        # フェーズ構成
        add_styled_heading(doc, "フェーズ構成", level=3)
        phase_data = [[ph["type"], f'{ph["start"]}〜{ph["end"]}', ph.get("note", "")] for ph in prod.get("phases", [])]
        if phase_data:
            add_table(doc, ["フェーズ", "時間帯", "内容"], phase_data, col_widths=[3, 3.5, 16])

        # テクニック実例（引用付き）
        start_min = _time_to_min(prod["start_time"])
        end_min = _time_to_min(prod["end_time"])

        add_styled_heading(doc, "テクニック実例（原文引用）", level=3)
        for tech_key in ["scarcity", "social_proof", "anchoring", "comparison", "authority", "urgency"]:
            examples = []
            for chunk in chunks:
                techs = chunk.get("techniques", {})
                for item in techs.get(tech_key, []):
                    ts_str = item.get("timestamp", "")
                    try:
                        ts_min = int(ts_str.split(":")[0]) if ":" in ts_str else 0
                    except:
                        ts_min = 0
                    if start_min - 5 <= ts_min <= end_min + 10:
                        examples.append(item)

            if not examples:
                continue

            p = doc.add_paragraph()
            run = p.add_run(f"■ {tech_jp[tech_key]}（{len(examples)}件）")
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(0x4e, 0x79, 0xa7)

            for ex in examples[:6]:
                quote = ex.get("quote", ex.get("text", ""))
                ts_str = ex.get("timestamp", "")
                context = ex.get("context", "")
                text = f"[{ts_str}] {quote}"
                if context:
                    text += f" — {context}"
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                run = p.add_run(text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                p.paragraph_format.space_after = Pt(1)

    # ============================================================
    # 7. Gemini 視覚分析ハイライト
    # ============================================================
    if gemini:
        doc.add_page_break()
        add_styled_heading(doc, "7. 【詳細】Gemini 視覚分析ハイライト", level=1)
        add_body(doc, "60枚のスクリーンショットをGemini 2.5 Flashで分析。テロップ・価格表示・QRコード・画面レイアウト等の視覚要素を抽出。以下は各バッチの分析結果。", size=9)

        batch_num = 0
        for r in gemini.get("results", []):
            analysis_text = r.get("gemini_analysis", "")
            if len(analysis_text) > 100:
                batch_num += 1
                add_styled_heading(doc, f"バッチ {batch_num}: {r['timestamp_display']} — {r['reason']}", level=3)
                p = doc.add_paragraph()
                run = p.add_run(f"スクリーンショット: {r.get('screenshot', '')} | {r.get('description', '')}")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                # 分析テキスト（最大1200文字）
                text = analysis_text[:1200]
                if len(analysis_text) > 1200:
                    text += "\n... (以下省略)"
                add_body(doc, text, size=8)

                if batch_num >= 8:
                    break

    # ============================================================
    # フッター
    # ============================================================
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Claude Code テキスト分析 + Gemini 2.5 Flash 視覚分析\n")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run = p.add_run(f"生成日: {info['analysis_date']}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 保存
    docx_path = os.path.join(output_dir, "analysis_report.docx")
    doc.save(docx_path)
    print(f"DOCX出力: {docx_path}")

    # PDF変換
    # LibreOffice or pandoc が使えればPDF変換
    pdf_path = os.path.join(output_dir, "analysis_report.pdf")

    # Try LibreOffice
    lo_result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        capture_output=True, text=True,
    )
    if lo_result.returncode == 0 and os.path.exists(pdf_path):
        print(f"PDF出力: {pdf_path}")
    else:
        # Try docx2pdf
        try:
            subprocess.run(
                [sys.executable, "-m", "pip", "install", "docx2pdf", "-q"],
                capture_output=True,
            )
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            print(f"PDF出力: {pdf_path}")
        except Exception as e:
            print(f"PDF変換スキップ（{e}）。DOCXファイルを直接お使いください: {docx_path}")


def _time_to_min(t):
    parts = t.split(":")
    if len(parts) == 2:
        return int(parts[0]) + int(parts[1]) / 60
    return 0


if __name__ == "__main__":
    main()
